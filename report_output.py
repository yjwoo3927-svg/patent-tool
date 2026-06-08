"""
특허 서지정보 → Word 보고서 출력 모듈
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import Counter
from datetime import datetime
import os


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def create_patent_report(patents: list, keyword: str, assignee: str,
                         output_path: str) -> str:
    doc = Document()

    # 페이지 마진
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # 스타일 기본값
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # ── 표지 ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run("특허 · 선행기술 서지정보 조사 보고서")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"검색어: {keyword or '(전체)'}  |  출원인: {assignee or '(전체)'}\n"
        f"수집일: {datetime.now().strftime('%Y년 %m월 %d일')}"
    )
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

    doc.add_paragraph()
    divider = doc.add_paragraph("─" * 55)
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ── 1. 조사 개요 ──────────────────────────────────────────────────────
    h1 = doc.add_heading("1. 조사 개요", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("조사 목적", "경쟁사 특허 모니터링 및 기술 분야 선행기술 조사"),
        ("검색 키워드", keyword or "(조건 없음)"),
        ("출원인/권리자", assignee or "(조건 없음)"),
        ("검색 데이터베이스", "KIPRIS (한국), USPTO (미국), EPO (유럽), WIPO (국제), CNIPA (중국)"),
        ("수집 특허 수", f"총 {len(patents)}건"),
    ]
    for i, (label, val) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = val
        _set_cell_bg(row.cells[0], "E8EDF2")
        _set_cell_border(row.cells[0])
        _set_cell_border(row.cells[1])

    doc.add_paragraph()

    # ── 2. 수집 결과 요약 ─────────────────────────────────────────────────
    h2 = doc.add_heading("2. 수집 결과 요약", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    country_cnt = Counter(p["country"] for p in patents)
    assignee_cnt = Counter(p["assignee"] for p in patents)

    # 국가별 현황 표
    doc.add_heading("2.1 국가별 수집 현황", level=2)
    tbl = doc.add_table(rows=1 + len(country_cnt) + 1, cols=3)
    tbl.style = "Table Grid"
    headers = ["국가", "건수", "비율(%)"]
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        _set_cell_bg(cell, "1A3A5C")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_border(cell)

    total = len(patents)
    for i, (country, cnt) in enumerate(sorted(country_cnt.items(), key=lambda x: -x[1]), 1):
        row = tbl.rows[i]
        row.cells[0].text = country
        row.cells[1].text = str(cnt)
        row.cells[2].text = f"{cnt/total*100:.1f}%"
        if i % 2 == 0:
            _set_cell_bg(row.cells[0], "F7F9FC")
            _set_cell_bg(row.cells[1], "F7F9FC")
            _set_cell_bg(row.cells[2], "F7F9FC")
        for cell in row.cells:
            _set_cell_border(cell)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 합계 행
    total_row = tbl.rows[-1]
    total_row.cells[0].text = "합계"
    total_row.cells[0].paragraphs[0].runs[0].font.bold = True
    total_row.cells[1].text = str(total)
    total_row.cells[1].paragraphs[0].runs[0].font.bold = True
    total_row.cells[2].text = "100%"
    for cell in total_row.cells:
        _set_cell_bg(cell, "E8EDF2")
        _set_cell_border(cell)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 주요 출원인 표
    doc.add_heading("2.2 주요 출원인/권리자 현황 (Top 10)", level=2)
    top_n = min(10, len(assignee_cnt))
    tbl2 = doc.add_table(rows=1 + top_n, cols=3)
    tbl2.style = "Table Grid"
    for j, h in enumerate(["순위", "출원인/권리자", "건수"]):
        cell = tbl2.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        _set_cell_bg(cell, "1A3A5C")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_border(cell)

    for i, (asgn, cnt) in enumerate(assignee_cnt.most_common(top_n), 1):
        row = tbl2.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = asgn
        row.cells[2].text = str(cnt)
        if i % 2 == 0:
            for cell in row.cells:
                _set_cell_bg(cell, "F7F9FC")
        for cell in row.cells:
            _set_cell_border(cell)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── 3. 특허별 상세 서지정보 ───────────────────────────────────────────
    h3 = doc.add_heading("3. 특허별 상세 서지정보", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    for i, p in enumerate(patents, 1):
        num_p = doc.add_paragraph()
        num_run = num_p.add_run(f"[{i}] {p.get('patent_number', '')}")
        num_run.font.bold = True
        num_run.font.size = Pt(11)
        num_run.font.color.rgb = RGBColor(0x1D, 0x9E, 0x75)

        detail_tbl = doc.add_table(rows=7, cols=2)
        detail_tbl.style = "Table Grid"
        detail_data = [
            ("발명의 명칭", p.get("title", "")),
            ("출원인/권리자", p.get("assignee", "")),
            ("발명자", p.get("inventors", "")),
            ("IPC 코드", p.get("ipc_codes", "")),
            ("출원일 / 공개일", f"{p.get('application_date', '')} / {p.get('publication_date', '')}"),
            ("국가 / 현황", f"{p.get('country', '')} / {p.get('status', '')}"),
            ("초록", p.get("abstract", "")),
        ]
        for j, (label, val) in enumerate(detail_data):
            row = detail_tbl.rows[j]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            row.cells[1].text = val
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
            _set_cell_bg(row.cells[0], "EEF2F7")
            _set_cell_border(row.cells[0])
            _set_cell_border(row.cells[1])

        detail_tbl.columns[0].width = Cm(3.5)
        detail_tbl.columns[1].width = Cm(12.5)
        doc.add_paragraph()

        if i < len(patents) and i % 3 == 0:
            doc.add_page_break()

    # ── 4. 분석 의견 ──────────────────────────────────────────────────────
    doc.add_page_break()
    h4 = doc.add_heading("4. 조사 결과 분석 의견", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    ipc_cnt = Counter()
    for p in patents:
        for code in p.get("ipc_codes", "").split(";"):
            main = code.strip()[:5]
            if main:
                ipc_cnt[main] += 1

    opinion_items = [
        f"총 {len(patents)}건의 특허 서지정보를 수집하였습니다.",
        f"국가별로는 {', '.join(f'{c}: {n}건' for c, n in sorted(country_cnt.items(), key=lambda x: -x[1]))} 순으로 분포합니다.",
        f"출원인 중 가장 많은 특허를 보유한 기관은 {assignee_cnt.most_common(1)[0][0]} ({assignee_cnt.most_common(1)[0][1]}건)입니다." if assignee_cnt else "",
        f"주요 IPC 코드는 {', '.join(c for c, _ in ipc_cnt.most_common(3))} 등으로, 해당 기술 분야의 핵심 특허군이 형성되어 있습니다." if ipc_cnt else "",
        "수집된 특허에 대해 청구항 분석 및 회피 설계 검토를 권고합니다.",
    ]
    for item in opinion_items:
        if item:
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(item).font.size = Pt(10)

    doc.save(output_path)
    return output_path
